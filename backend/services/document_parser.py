"""文档解析与滑动窗口分块"""
import csv
import json
import logging
import os
import re
from io import BytesIO, StringIO

import chardet
import fitz
from docx import Document as DocxDocument
import openpyxl

logger = logging.getLogger("xin-ai.document_parser")

CHUNK_SIZE = 512
CHUNK_OVERLAP = 64


def _normalize_qa_items(raw_items: list[dict]) -> list[dict]:
    normalized: list[dict] = []
    for item in raw_items:
        question = str(item.get("question", "")).strip()
        answer = str(item.get("answer", "")).strip()
        if not question or not answer:
            continue
        if answer == "答案正在整理中":
            continue
        normalized.append({
            "id": len(normalized) + 1,
            "question": question,
            "answer": answer,
        })
    return normalized


def _parse_json_content(content: bytes) -> list[dict]:
    try:
        payload = json.loads(content.decode("utf-8-sig"))
    except json.JSONDecodeError as exc:
        raise ValueError(f"JSON 解析失败：{exc}") from exc

    if isinstance(payload, dict) and isinstance(payload.get("data"), list):
        payload = payload["data"]
    if not isinstance(payload, list):
        raise ValueError("JSON 文件格式错误：应为数组，元素包含 question / answer 字段。")

    items = [item for item in payload if isinstance(item, dict)]
    return _normalize_qa_items(items)


def _parse_csv_content(content: bytes) -> list[dict]:
    try:
        text = content.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ValueError("CSV 编码不支持，请使用 UTF-8 编码。") from exc

    reader = csv.DictReader(StringIO(text))
    if not reader.fieldnames:
        raise ValueError("CSV 文件缺少表头。")

    fieldnames = [f.strip() for f in reader.fieldnames if f]
    q_key = next((f for f in fieldnames if "问题" in f or f.lower() in {"question", "q"}), "")
    a_key = next((f for f in fieldnames if "答案" in f or f.lower() in {"answer", "a"}), "")
    if not q_key or not a_key:
        raise ValueError("CSV 表头需包含“问题/答案”字段（或 question/answer）。")

    rows: list[dict] = []
    for row in reader:
        rows.append({
            "question": row.get(q_key, ""),
            "answer": row.get(a_key, ""),
        })
    return _normalize_qa_items(rows)


def _parse_xlsx_sheet(ws) -> list[dict]:
    rows: list[dict] = []

    if ws.title == "精选问题":
        for row in ws.iter_rows(min_row=3, max_row=ws.max_row, values_only=True):
            if not row:
                continue
            question = row[1] if len(row) > 1 else ""
            answer = row[2] if len(row) > 2 else ""
            rows.append({"question": question, "answer": answer})
        parsed = _normalize_qa_items(rows)
        if parsed:
            return parsed

    header_row_idx = -1
    q_col = -1
    a_col = -1
    for idx, row in enumerate(
        ws.iter_rows(min_row=1, max_row=min(15, ws.max_row), values_only=True),
        start=1,
    ):
        if not row:
            continue
        cells = [str(cell).strip() if cell is not None else "" for cell in row]
        if q_col < 0:
            q_col = next((i for i, c in enumerate(cells) if "问题" in c or c.lower() == "question"), -1)
        if a_col < 0:
            a_col = next((i for i, c in enumerate(cells) if "答案" in c or c.lower() == "answer"), -1)
        if q_col >= 0 and a_col >= 0:
            header_row_idx = idx
            break

    if header_row_idx < 0:
        return []

    for row in ws.iter_rows(min_row=header_row_idx + 1, max_row=ws.max_row, values_only=True):
        if not row:
            continue
        question = row[q_col] if len(row) > q_col else ""
        answer = row[a_col] if len(row) > a_col else ""
        rows.append({"question": question, "answer": answer})
    return _normalize_qa_items(rows)


def _parse_xlsx_content(content: bytes) -> list[dict]:
    try:
        wb = openpyxl.load_workbook(BytesIO(content), data_only=True)
    except Exception as exc:
        raise ValueError(f"Excel 解析失败：{exc}") from exc

    if "精选问题" in wb.sheetnames:
        parsed = _parse_xlsx_sheet(wb["精选问题"])
        if parsed:
            return parsed

    for name in wb.sheetnames:
        parsed = _parse_xlsx_sheet(wb[name])
        if parsed:
            return parsed

    return []


def _parse_pdf(content: bytes) -> list[tuple[str, int]]:
    out: list[tuple[str, int]] = []
    try:
        doc = fitz.open(stream=content, filetype="pdf")
        for page_idx in range(len(doc)):
            page = doc[page_idx]
            text = page.get_text() or ""
            if text.strip():
                out.append((text.strip(), page_idx + 1))
        doc.close()
    except Exception as e:
        logger.warning("PDF 解析异常: %s", e)
        raise
    return out


def _parse_docx(content: bytes) -> str:
    doc = DocxDocument(BytesIO(content))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _decode_text_bytes(content: bytes) -> str:
    det = chardet.detect(content)
    enc = (det.get("encoding") or "utf-8").lower()
    if enc in ("ascii", None):
        enc = "utf-8"
    try:
        return content.decode(enc)
    except UnicodeDecodeError:
        return content.decode("utf-8", errors="replace")


def _sliding_chunks(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    text = text.strip()
    if not text:
        return []
    if len(text) <= chunk_size:
        return [text]
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunks.append(text[start:end])
        if end >= len(text):
            break
        start = end - overlap
    return chunks


def _segments_to_chunks(
    segments: list[tuple[str, dict]],
    base_filename: str,
) -> list[dict]:
    results: list[dict] = []
    seg_idx = 0
    for text, meta in segments:
        page = meta.get("page")
        chunks = _sliding_chunks(text)
        for ci, chunk in enumerate(chunks):
            m = {
                "filename": base_filename,
                "page": page,
                "chunk_index": len(results),
                "segment_index": seg_idx,
                "sub_chunk_index": ci,
                **{k: v for k, v in meta.items() if k != "page"},
            }
            results.append({"text": chunk, "metadata": m})
        seg_idx += 1
    return results


def parse_document(filename: str, content: bytes) -> list[dict]:
    if not filename or not content:
        return []
    ext = os.path.splitext(filename)[1].lower()
    base = os.path.basename(filename)
    segments: list[tuple[str, dict]] = []

    try:
        if ext == ".pdf":
            for text, page in _parse_pdf(content):
                segments.append((text, {"page": page}))
        elif ext == ".docx":
            text = _parse_docx(content)
            if text:
                segments.append((text, {"page": None}))
        elif ext in (".txt", ".md"):
            text = _decode_text_bytes(content)
            if text.strip():
                segments.append((text.strip(), {"page": None}))
        elif ext == ".json":
            qa_list = _parse_json_content(content)
            for item in qa_list:
                q = item["question"]
                a = item["answer"]
                block = f"问题：{q}\n答案：{a}"
                segments.append((block, {"page": None, "qa_id": item.get("id"), "format": "qa_json"}))
        elif ext == ".csv":
            qa_list = _parse_csv_content(content)
            for item in qa_list:
                q = item["question"]
                a = item["answer"]
                block = f"问题：{q}\n答案：{a}"
                segments.append((block, {"page": None, "format": "qa_csv"}))
        elif ext == ".xlsx":
            qa_list = _parse_xlsx_content(content)
            if not qa_list:
                raise ValueError("未解析到有效问答，请检查 Excel 模板。")
            for item in qa_list:
                q = item["question"]
                a = item["answer"]
                block = f"问题：{q}\n答案：{a}"
                segments.append((block, {"page": None, "format": "qa_xlsx"}))
        else:
            raise ValueError(f"不支持的文件类型: {ext}")
    except ValueError:
        raise
    except Exception as e:
        logger.exception("parse_document 失败: %s", e)
        raise ValueError(str(e)) from e

    return _segments_to_chunks(segments, base)
